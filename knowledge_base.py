# -*- coding: utf-8 -*-
"""Local knowledge-base indexing and retrieval helpers.

This module intentionally avoids external APIs. It builds a SQLite inverted
index over normalized Chinese/English terms and ranks chunks locally.
"""

import math
import os
import re
import json
import sqlite3
from collections import Counter, defaultdict
from datetime import datetime


SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.txt', '.md'}

S2T = str.maketrans({
    '\u4e1a': '\u696d', '\u52a1': '\u52d9', '\u4ef7': '\u50f9', '\u62a5': '\u5831',
    '\u56e2': '\u5718', '\u8d44': '\u8cc7', '\u76d6': '\u84cb', '\u9669': '\u96aa',
    '\u79cd': '\u7a2e', '\u8d23': '\u8cac', '\u8d54': '\u8ce0', '\u73b0': '\u73fe',
    '\u8d26': '\u8cec', '\u6237': '\u6236', '\u6761': '\u689d', '\u4ea7': '\u7522',
    '\u533b': '\u91ab', '\u5bff': '\u58fd', '\u8f7b': '\u8f15', '\u7f34': '\u7e73',
    '\u8d39': '\u8cbb', '\u6da6': '\u6f64', '\u7a0e': '\u7a05', '\u540e': '\u5f8c',
    '\u50a8': '\u5132', '\u5907': '\u5099', '\u603b': '\u7e3d', '\u5f00': '\u958b',
    '\u8425': '\u71df', '\u8fd0': '\u904b', '\u8d22': '\u8ca1', '\u5ba1': '\u5be9',
    '\u8ba1': '\u8a08', '\u4e1c': '\u6771', '\u5458': '\u54e1', '\u5185': '\u5167',
    '\u957f': '\u9577', '\u573a': '\u5834', '\u7ea2': '\u7d05', '\u4e07': '\u842c',
    '\u989d': '\u984d', '\u72b9': '\u7336', '\u7ed9': '\u7d66', '\u79ef': '\u7a4d',
    '\u5151': '\u514c', '\u8d60': '\u8d08'
})

T2S = str.maketrans({v: k for k, v in S2T.items()})

SYNONYMS = {
    '\u6350\u6b3e': ['\u6350\u8d60', '\u6350\u8d08', '\u6148\u5584', '\u516c\u76ca'],
    '\u6350\u8d60': ['\u6350\u6b3e', '\u6350\u8d08', '\u6148\u5584', '\u516c\u76ca'],
    '\u6350\u8d08': ['\u6350\u6b3e', '\u6350\u8d60', '\u6148\u5584', '\u516c\u76ca'],
    '\u793e\u533a': ['\u793e\u5340', '\u793e\u6703', '\u793e\u4f1a'],
    '\u793e\u5340': ['\u793e\u533a', '\u793e\u6703', '\u793e\u4f1a'],
    '\u516c\u76ca': ['\u6148\u5584', '\u6350\u6b3e', '\u6350\u8d08'],
    '\u6148\u5584': ['\u516c\u76ca', '\u6350\u6b3e', '\u6350\u8d08'],
}

BILINGUAL_CONCEPTS = {
    'donation': ['捐款', '捐赠', '慈善', '公益'],
    'donations': ['捐款', '捐赠', '慈善', '公益'],
    'charitable': ['慈善', '公益', '捐款'],
    'charity': ['慈善', '公益', '捐款'],
    'community': ['社区', '社會', '社会', '公益'],
    'valueofnewbusiness': ['新业务价值', '新業務價值', '新业务', 'VONB'],
    'vnb': ['新业务价值', '新業務價值', 'value of new business'],
    'newbusinessvalue': ['新业务价值', '新業務價值'],
    'annualizednewpremiums': ['年化新保费', '年化新保費', 'ANP'],
    'anp': ['年化新保费', '年化新保費', 'annualized new premiums'],
    'coverratio': ['覆盖率', '覆蓋率', '资本覆盖率', '資本覆蓋率'],
    'coverage ratio': ['覆盖率', '覆蓋率', '资本覆盖率', '資本覆蓋率'],
    'capitalcoverratio': ['资本覆盖率', '資本覆蓋率', '集团资本覆盖率'],
    'groupcapital': ['集团资本', '集團資本'],
    'solvency': ['偿付能力', '償付能力', '资本充足'],
    'embeddedvalue': ['内含价值', '內涵價值', 'EV'],
    'operatingprofit': ['营运利润', '營運溢利', '经营利润'],
    'aftertax': ['税后', '稅後'],
    'premium': ['保费', '保費'],
    'premiums': ['保费', '保費'],
    'policy': ['保单', '保單'],
    'policies': ['保单', '保單'],
    'insured': ['被保险人', '受保人', '受保人'],
    'beneficiary': ['受益人'],
    'claim': ['理赔', '理賠', '赔付', '賠付'],
    'claims': ['理赔', '理賠', '赔付', '賠付'],
    'exclusion': ['免责', '责任免除', '責任免除'],
    'waitingperiod': ['等待期'],
    'cashvalue': ['现金价值', '現金價值'],
    'criticalillness': ['重疾', '重大疾病'],
    'medical': ['医疗险', '醫療險', '医疗'],
    'lifeinsurance': ['寿险', '壽險', '人寿保险'],
    'annualreport': ['年报', '年報', '年度报告'],
}

for key, values in list(BILINGUAL_CONCEPTS.items()):
    compact_key = normalize_key = re.sub(r'\s+', '', key.lower())
    if compact_key != key:
        BILINGUAL_CONCEPTS.setdefault(compact_key, values)

STOP_TERMS = {
    '\u8fd9\u4e2a', '\u90a3\u4e2a', '\u4ec0\u4e48', '\u600e\u4e48', '\u591a\u5c11',
    '\u662f\u5426', '\u53ef\u4ee5', '\u4e00\u4e0b', '\u5e2e\u6211', '\u8bf7\u95ee',
    '\u8d44\u6599', '\u6587\u4ef6', '\u91cc\u9762', '\u5173\u4e8e', '\u5982\u679c',
    '\u9700\u8981', '\u67e5\u8be2', '\u544a\u8bc9', '\u6211\u4eec', '\u4f60\u4eec',
    '\u6709\u6ca1\u6709', '\u662f\u4e0d\u662f', '\u4e3a\u4ec0\u4e48',
    '\u516c\u53f8', '\u96c6\u56e2', '\u672c\u96c6\u5718', '\u53bb\u5e74', '\u4eca\u5e74',
    'the', 'and', 'or', 'of', 'in', 'on', 'for', 'to', 'a', 'an', 'is', 'are',
    'new', 'business', 'value', 'amount', 'report',
}

NUMERIC_QUESTION_TERMS = {'\u591a\u5c11', '\u91d1\u989d', '\u6bd4\u4f8b', '\u51e0', '\u7387', 'amount', 'rate'}
NUMERIC_PATTERN = re.compile(r'(\d[\d,]*(?:\.\d+)?\s*(?:%|\u842c|\u4e07|\u5104|\u7f8e\u5143|\u6e2f\u5143|元|美元)?)')
YEAR_TERM_PATTERN = re.compile(r'^(?:19|20)\d{2}年?$')
SIMPLIFIED_HINT_CHARS = set('业东丝丢两严丧个丰临为丽举么义乌乐乔习乡书买乱争于亏云亚产亩亲亵亿仅从仑仓仪们价众优伙会伛伞伟传伤伦伪伫体余佣佥侠侣侥侦侧侨侩侪侬俣俦俨俩俪俭债倾偬偻偾偿傥傧储傩儿兑党兰关兴养兽冁内冈册写军农冯冲决况冻净凄准凉减凑凛凤凭凯击凿刍刘则刚创删别刬刭刹刿剀剂剐剑剥剧劝办务动励劲劳势勋匀匦匮区医华协单卖卢卤卧卫却卺厂厅历厉压厌厍厕厘厢厣厦厨厩县参双发变叙叠叶号叹叽吓吕吗吣吨听启吴呐呒呓呕呖呗员呙呛呜咏咙咛咝响哑哒哓哔哕哗哙哜哝哟唛唝唠唡唢唤啧啬啭啮啰啴啸喷喽喾嗫嗳嘘嘤嘱噜噼嚣团园囱围囵国图圆圣圹场坏块坚坛坜坝坞坟坠垄垅垆垒垦垩垫垭垯垱垲垴埘埙埚埯堑堕墙壮声壳壶处备复够头夹夺奁奂奋奖奥妆妇妈妩妪妫姗姜娄娅娆娇娈娱娲娴婳婴婵婶媪嫒嫔嫱嬷孙学孪宁宝实宠审宪宫宽宾寝对寻导寿将尔尘尝尧尸尽层屃屉届属屡屦屿岁岂岖岗岘岚岛岭岳岽岿峃峄峡峣峤峥峦崂崃崭嵘嵚嵝巅巩巯币帅师帐帘帜带帧帮帱帻帼幂庄庆庐庑库应庙庞废庼廪开异弃张弥弪弯弹强归当录彦彻径徕忆忏忧忾怀态怂怃怄怅怆怜总怼怿恋恳恶恸恹恺恻恼恽悦悬悭悯惊惧惨惩惫惬惭惮惯愠愤愦愿慑懑懒懔戆戋戏戗战戬户扎扑托执扩扪扫扬扰抚抛抟抠抡抢护报担拟拢拣拥拦拧拨择挂挚挛挜挝挞挟挠挡挢挣挤挥挦捞损捡换捣据捻掳掴掷掸掺揽揿搀搁搂搅携摄摅摆摇摈摊撄撑撵撷撸撺擞攒敌敛数斋斓斗斩断无旧时旷旸昙昼显晋晒晓晔晕晖暂术朴机杀杂权条来杨杩杰极构枞枢枣枪枫枭柜柠查栀栅标栈栉栊栋栌栎栏树栖样栾桊桠桡桢档桤桥桦桧桨桩梦梼梾检棂椁椟椠椢椤椭楼榄榅榇榈榉槚槛槟槠横樯樱橥橱橹橼檐檩欢欤欧欲歼殁殇残殒殓殚殡殴毁毂毕毙毡毵氇气氢氩氲汇汉汤汹沟没沣沤沥沦沧沨沩沪泞注泪泶泷泸泺泻泼泽泾洁洒洼浃浅浆浇浈浊测浍济浏浐浑浒浓浔浕涂涛涝涞涟涠涡涢涣涤润涧涨涩淀渊渌渍渎渐渑渔渖渗温湾湿溃溅溆溇滗滚滞滟滠满滢滤滥滦滨滩滪漤潆潇潋潍潜潴澜濑濒灏灭灯灵灾灿炀炉炖炜炝点炼炽烁烂烃烛烟烦烧烨烩烫烬热焕焖焘煴爱爷牍牦牵牺犊状犷犸犹狈狝狞独狭狮狯狰狱狲猃猎猕猡猪猫猬献獭玑玙玚玛玮环现玱玺珐珑珰珲琏琐琼瑶瑷璎瓒瓯电画畅畴疖疗疟疠疡疬疮疯疱疴痈痉痒痖痨痪痫痴瘅瘆瘗瘘瘪瘫瘾瘿癞癣癫皑皱皲盏盐监盖盗盘眍眦睁睐睑瞒瞩矫矶矾矿砀码砖砗砚砜砺砻础硁硕硖硗硙硚确硷碍碛碜碱礼祃祎祢祯祷祸禀禄禅离秃秆种积称秽秾稆税稣稳穑穷窃窍窑窜窝窥窦窭竖竞笃笋笔笕笺笼笾筑筚筛筜筝筹签简箓箦箧箨箩箪箫篑篓篮篱簖籁籴类籼粜粝粤粪粮糁糇紧絷纟纠纡红纣纤纥约级纨纩纪纫纬纭纯纰纱纲纳纵纶纷纸纹纺纽纾线绀绁绂练组绅细织终绉绊绍绎经绑绒结绕绗绘给绚绛络绝绞统绠绡绢绣绥绦继绩绪绫续绮绯绰绱绲绳维绵绶绷绸绹绺综绽绾绿缀缁缂缃缄缅缆缇缈缉缊缋缌缍缎缏缐缑缒缓缔缕编缗缘缙缚缛缜缝缟缠缡缢缣缤缥缦缧缨缩缪缫缬缭缮缯缰缱缲缳缴缵罂网罗罚罢罴羁羟羡翘耢耧耸耻聂聋职聍联聩聪肃肠肤肮肴肾肿胀胁胆胜胧胨胪胫胶脉脍脏脐脑脓脔脚脱脶脸腊腌腘腭腻腼腽腾膑臜舆舣舰舱舻艰艳艺节芈芗芜芦苁苇苈苋苌苍苎苏苘苹茎茏茑茔茕茧荆荐荙荚荛荜荞荟荠荡荣荤荥荦荧荨荩荪荫荬荭荮药莅莱莲莳莴莶获莸莹莺莼萝萤营萦萧萨葱蒇蒉蒋蒌蓝蓟蓠蓣蓥蓦蔷蔹蔺蔼蕰蕲蕴薮藓虏虑虚虫虬虮虽虾虿蚀蚁蚂蚕蚝蚬蛊蛎蛏蛮蛰蛱蛲蛳蛴蜕蜗蜡蝇蝈蝉蝎蝼蝾螀螨蟏衅衔补衬衮袄袅袆袜袭袯装裆裈裢裣裤裥褛褴襁襕见观觃规觅视觇览觉觊觋觌觎觏觐觑觞触觯訚詟誉誊讠计订讣认讥讦讧讨让讪讫训议讯记讲讳讴讵讶讷许讹论讻讼讽设访诀证诂诃评诅识诈诉诊诋诌词诎诏译诒诓诔试诖诗诘诙诚诛诜话诞诟诠诡询诣诤该详诧诨诩诫诬语诮误诰诱诲诳说诵诶请诸诹诺读诼诽课诿谀谁谂调谄谅谆谈谊谋谌谍谎谏谐谑谒谓谔谕谖谗谘谙谚谛谜谝谟谠谡谢谣谤谥谦谧谨谩谪谫谬谭谮谯谰谱谲谳谴谵谷豮贝贞负贡财责贤败账货质贩贪贫贬购贮贯贰贱贲贳贴贵贶贷贸费贺贻贼贽贾贿赀赁赂赃资赅赆赇赈赉赊赋赌赍赎赏赐赑赒赓赔赖赗赘赙赚赛赜赝赞赠赡赢赣赵赶趋趱跃跄跞践跶跷跸跹跻踊踌踪踬踯蹑蹒蹰蹿躏躜躯车轧轨轩轫转轮软轭轰轱轲轳轴轵轶轷轸轹轺轻轼载轾轿辀辁辂较辄辅辆辇辈辉辊辋辌辍辎辏辐辑输辔辕辖辗辘辙辚辞辩辫边辽达迁过迈运还这进远违连迟迩迳迹适选逊递逦逻遗遥邓邝邬邮邹邺邻郁郏郐郑郓郦郧郸酝酦酱酽酾酿释里鉴銮錾钅针钉钊钋钌钍钏钐钒钓钔钕钗钘钙钚钛钜钝钞钟钠钡钢钣钤钥钦钧钨钩钪钫钬钭钮钯钰钱钲钳钴钵钶钷钹钺钻钼钽钾铀铁铂铃铄铅铆铈铉铊铋铌铍铎铐铑铒铕铖铗铙铛铜铝铞铟铠铡铢铣铤铥铧铨铩铪铫铭铬铭铮铯铰铱铲铳铴铵银铷铸铹铺铻铼铽链铿销锁锂锃锄锅锆锇锋锌锏锐锑锒锓锔锕锖锗错锚锛锜锝锞锟锡锢锣锤锥锦锨锩锪锫锬锭键锯锰锱锲锴锵锶锷锸锹锺锻锼锽锾锿镀镁镂镃镄镅镇镉镊镌镍镎镏镐镑镒镓镔镕镖镗镘镙镚镜镝镞镟镠镡镢镣镤镥镦镧镨镩镪镫镬镭镯镰镱镲镳镴镵长门闩闪闫闭问闯闰闱闲闳间闵闶闷闸闹闺闻闼闽闾阀阁阂阃阄阅阆阈阉阊阋阌阍阎阏阐阑阒阔阕阖阗阙阚队阳阴阵阶际陆陇陈陉陕陧陨险随隐隶隽难雏雠雳雾霁霉霭靓静面靥鞑鞒鞯韦韧韩韪韫韬韵页顶顷项顺须顼顽顾顿颀颁颂预颅领颇颈颉颊颋颌颍颎颏颐频颓颖颗题颚颜额颞颟颠颡颢颤颥颦颧风飏飐飒飓飔飕飖飘飙飚飞饣饥饧饨饩饪饫饬饭饮饯饰饱饲饳饴饵饶饷饸饹饺饼饽饿馀馁馂馄馅馆馈馊馋馍馎馏馐馑馒馓馔馕马驭驮驯驰驱驳驴驶驷驸驹驻驼驽驾驿骀骁骂骄骅骆骇骈骊骋验骏骐骑骒骓骖骗骘骚骛骜骝骞骟骠骡骢骣骤骥骦骧髅髋髌鬓魇魉鱼鱿鲁鲂鲅鲆鲇鲈鲋鲍鲎鲐鲑鲒鲔鲕鲚鲛鲜鲞鲟鲠鲡鲢鲣鲤鲥鲦鲧鲨鲩鲫鲭鲮鲰鲱鲲鲳鲴鲵鲶鲷鲸鲺鲻鲼鲽鲾鲿鳀鳁鳃鳄鳅鳆鳇鳈鳉鳊鳋鳌鳍鳎鳏鳐鳓鳔鳕鳖鳗鳘鳙鳜鳝鳟鳢鳣鸟鸠鸡鸢鸣鸥鸦鸧鸨鸩鸪鸫鸬鸭鸯鸱鸲鸳鸵鸶鸷鸸鸹鸺鸻鸽鸾鸿鹁鹂鹃鹄鹅鹆鹇鹈鹉鹊鹋鹏鹑鹕鹗鹘鹚鹛鹜鹞鹣鹤鹦鹧鹨鹩鹪鹫鹬鹭鹯鹰鹱鹳鹴鹾麦黄黉黩黪黾鼋鼍鼹齐齑齿龀龁龂龃龄龅龆龇龈龉龊龋龌龙龚龛龟')
TRADITIONAL_HINT_CHARS = {char.translate(S2T) for char in SIMPLIFIED_HINT_CHARS if char.translate(S2T) != char}


def resolve_path(base_dir, path):
    value = str(path or '')
    if not os.path.isabs(value):
        value = os.path.join(base_dir, value)
    return value


def connect(db_path):
    os.makedirs(os.path.dirname(db_path), exist_ok=True)
    conn = sqlite3.connect(db_path, timeout=15)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def init_db(db_path):
    with connect(db_path) as conn:
        conn.executescript("""
CREATE TABLE IF NOT EXISTS kb_documents (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    collection TEXT NOT NULL DEFAULT '保险知识',
    title TEXT NOT NULL,
    source_filename TEXT,
    stored_filename TEXT,
    source_type TEXT,
    trusted_level TEXT DEFAULT '正式资料',
    product_name TEXT,
    version TEXT,
    effective_date TEXT,
    enabled INTEGER DEFAULT 1,
    chunk_count INTEGER DEFAULT 0,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS kb_chunks (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    document_id INTEGER NOT NULL,
    chunk_index INTEGER NOT NULL,
    title TEXT,
    content TEXT NOT NULL,
    page_number INTEGER,
    section_path TEXT,
    enabled INTEGER DEFAULT 1,
    created_at TEXT NOT NULL,
    FOREIGN KEY(document_id) REFERENCES kb_documents(id) ON DELETE CASCADE
);
CREATE TABLE IF NOT EXISTS kb_chunk_terms (
    term TEXT NOT NULL,
    chunk_id INTEGER NOT NULL,
    document_id INTEGER NOT NULL,
    tf REAL NOT NULL DEFAULT 1,
    PRIMARY KEY(term, chunk_id),
    FOREIGN KEY(chunk_id) REFERENCES kb_chunks(id) ON DELETE CASCADE,
    FOREIGN KEY(document_id) REFERENCES kb_documents(id) ON DELETE CASCADE
);
CREATE INDEX IF NOT EXISTS idx_kb_documents_collection ON kb_documents(collection, enabled);
CREATE INDEX IF NOT EXISTS idx_kb_chunks_document ON kb_chunks(document_id, chunk_index);
CREATE INDEX IF NOT EXISTS idx_kb_chunk_terms_term ON kb_chunk_terms(term);
CREATE INDEX IF NOT EXISTS idx_kb_chunk_terms_chunk ON kb_chunk_terms(chunk_id);
CREATE INDEX IF NOT EXISTS idx_kb_chunk_terms_document ON kb_chunk_terms(document_id);
""")
        ensure_document_analysis_columns(conn)


def ensure_document_analysis_columns(conn):
    existing = {
        row['name']
        for row in conn.execute("PRAGMA table_info(kb_documents)").fetchall()
    }
    columns = {
        'dominant_language': "TEXT DEFAULT ''",
        'language_summary': "TEXT DEFAULT ''",
        'language_stats': "TEXT DEFAULT ''",
        'content_char_count': "INTEGER DEFAULT 0",
        'english_word_count': "INTEGER DEFAULT 0",
        'simplified_char_count': "INTEGER DEFAULT 0",
        'traditional_char_count': "INTEGER DEFAULT 0",
        'translated_term_count': "INTEGER DEFAULT 0",
    }
    for column, definition in columns.items():
        if column not in existing:
            conn.execute(f"ALTER TABLE kb_documents ADD COLUMN {column} {definition}")


def normalize_text(text):
    text = text or ''
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def normalize_for_search(text):
    return re.sub(r'\s+', '', str(text or '').lower())


def analyze_text_language(text):
    raw = str(text or '')
    total_chars = len(raw)
    chinese_chars = re.findall(r'[\u4e00-\u9fff]', raw)
    english_words = re.findall(r'[A-Za-z][A-Za-z\-]{1,}', raw)
    numbers = re.findall(r'\d[\d,]*(?:\.\d+)?', raw)
    simplified_count = sum(1 for char in chinese_chars if char in SIMPLIFIED_HINT_CHARS or char.translate(S2T) != char)
    traditional_count = sum(1 for char in chinese_chars if char in TRADITIONAL_HINT_CHARS or char.translate(T2S) != char)
    chinese_count = len(chinese_chars)
    english_count = len(english_words)
    concept_count = sum(1 for key in BILINGUAL_CONCEPTS if normalize_for_search(key) in normalize_for_search(raw))

    if chinese_count and english_count >= max(12, chinese_count // 8):
        dominant = '中英混合'
    elif chinese_count:
        if traditional_count > simplified_count * 1.4:
            dominant = '繁体中文'
        elif simplified_count > traditional_count * 1.4:
            dominant = '简体中文'
        else:
            dominant = '中文'
    elif english_count:
        dominant = '英文'
    else:
        dominant = '未知'

    stats = {
        'dominant_language': dominant,
        'content_char_count': total_chars,
        'chinese_char_count': chinese_count,
        'simplified_char_count': simplified_count,
        'traditional_char_count': traditional_count,
        'english_word_count': english_count,
        'number_count': len(numbers),
        'translated_term_count': concept_count,
    }
    summary_parts = [dominant, f'{total_chars}字']
    if chinese_count:
        summary_parts.append(f'中文{chinese_count}')
        if simplified_count or traditional_count:
            summary_parts.append(f'简{simplified_count}/繁{traditional_count}')
    if english_count:
        summary_parts.append(f'英文{english_count}词')
    if concept_count:
        summary_parts.append(f'概念映射{concept_count}')
    stats['language_summary'] = '，'.join(summary_parts)
    return stats


def update_document_language_profile(conn, document_id, text=None):
    if text is None:
        rows = conn.execute(
            "SELECT content FROM kb_chunks WHERE document_id = ? ORDER BY chunk_index",
            (document_id,)
        ).fetchall()
        text = '\n'.join(row['content'] or '' for row in rows)
    stats = analyze_text_language(text)
    conn.execute("""
        UPDATE kb_documents
        SET dominant_language = ?,
            language_summary = ?,
            language_stats = ?,
            content_char_count = ?,
            english_word_count = ?,
            simplified_char_count = ?,
            traditional_char_count = ?,
            translated_term_count = ?,
            updated_at = COALESCE(updated_at, ?)
        WHERE id = ?
    """, (
        stats['dominant_language'],
        stats['language_summary'],
        json.dumps(stats, ensure_ascii=False),
        stats['content_char_count'],
        stats['english_word_count'],
        stats['simplified_char_count'],
        stats['traditional_char_count'],
        stats['translated_term_count'],
        datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        document_id,
    ))
    return stats


def text_variants(text):
    text = str(text or '').lower()
    variants = {text, text.translate(S2T), text.translate(T2S)}
    return {item for item in variants if item}


def add_term_with_variants(terms, term):
    if not term:
        return
    for variant in text_variants(term):
        compact_variant = normalize_for_search(variant)
        terms.add(variant)
        terms.add(compact_variant)
        for synonym in SYNONYMS.get(variant, []):
            terms.update(text_variants(synonym))
        for concept in BILINGUAL_CONCEPTS.get(compact_variant, []):
            terms.update(text_variants(normalize_for_search(concept)))


def bilingual_concept_terms_for_text(text):
    compact = normalize_for_search(text)
    terms = Counter()
    if not compact:
        return terms
    for key, values in BILINGUAL_CONCEPTS.items():
        compact_key = normalize_for_search(key)
        if compact_key and compact_key in compact:
            terms[compact_key] += 2.0
            for value in values:
                for variant in text_variants(normalize_for_search(value)):
                    terms[variant] += 2.5
        for value in values:
            compact_value = normalize_for_search(value)
            if compact_value and compact_value in compact:
                terms[compact_key] += 2.0
                for variant in text_variants(compact_value):
                    terms[variant] += 2.5
    return terms


def extract_relative_year_terms(query, now=None):
    now = now or datetime.now()
    mapping = {
        '\u4eca\u5e74': now.year,
        '\u53bb\u5e74': now.year - 1,
        '\u524d\u5e74': now.year - 2,
    }
    terms = []
    for marker, year in mapping.items():
        if marker in str(query or ''):
            terms.extend([str(year), f'{year}\u5e74'])
    return terms


def build_query_terms(query, limit=120):
    raw = str(query or '')
    terms = set()
    add_term_with_variants(terms, normalize_for_search(raw))
    for concept_term in bilingual_concept_terms_for_text(raw):
        add_term_with_variants(terms, concept_term)
    for year_term in extract_relative_year_terms(raw):
        add_term_with_variants(terms, year_term)
    for token in re.findall(r'[A-Za-z0-9_\-\u4e00-\u9fff]{2,}', raw.lower()):
        add_term_with_variants(terms, token)
    chinese = ''.join(re.findall(r'[\u4e00-\u9fff]', raw))
    for size in (2, 3, 4):
        for i in range(0, max(0, len(chinese) - size + 1)):
            add_term_with_variants(terms, chinese[i:i + size])
    useful = [term for term in terms if len(term) >= 2 and term not in STOP_TERMS]
    useful.sort(key=lambda item: (len(item), item), reverse=True)
    return useful[:limit]


def build_core_query_terms(query, limit=28):
    """Terms that should be covered by a strong answer candidate."""
    raw = str(query or '')
    query_variants = {normalize_for_search(item) for item in text_variants(raw)}
    concept_terms = set()
    for term in bilingual_concept_terms_for_text(raw):
        concept_terms.update(text_variants(normalize_for_search(term)))
    core = []
    for term in build_query_terms(raw, limit=220):
        if len(term) > 24 or term in STOP_TERMS:
            continue
        if any(stop in term for stop in STOP_TERMS if len(stop) >= 2 and re.search(r'[\u4e00-\u9fff]', stop)):
            continue
        in_original_query = any(term in variant for variant in query_variants)
        is_year_or_number = bool(YEAR_TERM_PATTERN.match(term) or re.fullmatch(r'\d[\d,.]*', term))
        is_translated_concept = term in concept_terms
        if in_original_query or is_year_or_number or is_translated_concept:
            core.append(term)
    core = sorted(set(core), key=lambda item: (query_term_weight(item, raw), len(item), item), reverse=True)
    return core[:limit]


def query_term_weight(term, query):
    term = normalize_for_search(term)
    if not term:
        return 0.0
    query_variants = {normalize_for_search(item) for item in text_variants(query)}
    weight = 1.0
    if any(term in variant for variant in query_variants):
        weight += 1.2
    if YEAR_TERM_PATTERN.match(term):
        weight += 3.2
    elif re.fullmatch(r'\d[\d,.]*', term):
        weight += 2.2
    if len(term) >= 6:
        weight += 2.0
    elif len(term) >= 4:
        weight += 1.2
    elif len(term) == 3:
        weight += 0.5
    if term in SYNONYMS:
        weight += 0.6
    return weight


def text_compact_variants(text):
    return {normalize_for_search(item) for item in text_variants(text)}


def term_in_compact_variants(term, compact_variants):
    term = normalize_for_search(term)
    return bool(term) and any(term in variant for variant in compact_variants)


def iter_chinese_segments(text):
    for segment in re.findall(r'[\u4e00-\u9fff]+', text or ''):
        if len(segment) >= 2:
            yield segment


def index_terms_for_text(text, base_weight=1.0):
    counts = Counter()
    raw_text = str(text or '').lower()
    variants = {raw_text, raw_text.translate(S2T)}
    for variant in variants:
        compact = normalize_for_search(variant)
        for token in re.findall(r'[A-Za-z0-9_\-]{2,}', variant):
            if token not in STOP_TERMS:
                counts[token] += base_weight * 1.2
        for token in re.findall(r'\d{4}\u5e74|\d{4}|\d[\d,]*(?:\.\d+)?', compact):
            counts[token] += base_weight * 1.5
        for segment in iter_chinese_segments(compact):
            if segment not in STOP_TERMS and len(segment) <= 12:
                counts[segment] += base_weight * 1.1
            for size in (2, 3):
                if len(segment) < size:
                    continue
                for i in range(0, len(segment) - size + 1):
                    term = segment[i:i + size]
                    if term not in STOP_TERMS:
                        counts[term] += base_weight
    for term, synonyms in SYNONYMS.items():
        if term in counts:
            for synonym in synonyms:
                for variant in text_variants(synonym):
                    counts[variant] += counts[term] * 0.7
    for term, value in bilingual_concept_terms_for_text(raw_text).items():
        counts[term] += value * base_weight
    return counts


def chunk_text(text, chunk_size=900, overlap=120):
    text = normalize_text(text)
    if not text:
        return []
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    chunks = []
    current = ''
    for paragraph in paragraphs:
        if len(current) + len(paragraph) + 2 <= chunk_size:
            current = f"{current}\n\n{paragraph}".strip()
            continue
        if current:
            chunks.append(current)
        if len(paragraph) <= chunk_size:
            current = paragraph
        else:
            start = 0
            step = max(1, chunk_size - overlap)
            while start < len(paragraph):
                chunks.append(paragraph[start:start + chunk_size].strip())
                start += step
            current = ''
    if current:
        chunks.append(current)
    return [chunk for chunk in chunks if chunk]


def extract_page_number(content):
    match = re.search(r'\[第(\d+)页\]', content or '')
    return int(match.group(1)) if match else None


def extract_text_from_docx(path):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，无法读取 Word 文件。请先安装 requirements.txt 中的依赖。") from exc
    doc = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                lines.append(' | '.join(cells))
    return '\n'.join(lines)


def extract_text_from_pdf(path):
    try:
        import fitz
    except ImportError:
        fitz = None
    if fitz:
        lines = []
        with fitz.open(path) as doc:
            for page_index, page in enumerate(doc, start=1):
                page_text = page.get_text("text").strip()
                if page_text:
                    lines.append(f"\n[第{page_index}页]\n{page_text}")
        return '\n'.join(lines)
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("缺少 PyMuPDF 或 pypdf，无法读取 PDF。请先安装 requirements.txt 中的依赖。") from exc
    reader = PdfReader(path)
    lines = []
    for page_index, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or '').strip()
        if page_text:
            lines.append(f"\n[第{page_index}页]\n{page_text}")
    return '\n'.join(lines)


def extract_text_from_plain(path):
    for encoding in ('utf-8', 'utf-8-sig', 'gbk', 'gb2312'):
        try:
            with open(path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def extract_text_from_xlsx(path):
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("缺少 openpyxl，无法读取 Excel 文件。请先安装 requirements.txt 中的依赖。") from exc
    workbook = load_workbook(path, read_only=True, data_only=True)
    lines = []
    for sheet in workbook.worksheets:
        lines.append(f"\n[工作表：{sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if values:
                lines.append(' | '.join(values))
    workbook.close()
    return '\n'.join(lines)


def extract_text_from_file(path, ext=None):
    ext = (ext or os.path.splitext(path)[1]).lower()
    if ext in ('.txt', '.md'):
        return extract_text_from_plain(path)
    if ext == '.docx':
        return extract_text_from_docx(path)
    if ext == '.pdf':
        return extract_text_from_pdf(path)
    if ext == '.xlsx':
        return extract_text_from_xlsx(path)
    raise RuntimeError(f"暂不支持 {ext} 文件。当前支持 PDF、DOCX、XLSX、TXT、MD。")


def build_index_for_document(conn, document_id):
    document = conn.execute("SELECT * FROM kb_documents WHERE id = ?", (document_id,)).fetchone()
    if not document:
        return 0
    if not document['language_summary']:
        update_document_language_profile(conn, document_id)
        document = conn.execute("SELECT * FROM kb_documents WHERE id = ?", (document_id,)).fetchone()
    conn.execute("DELETE FROM kb_chunk_terms WHERE document_id = ?", (document_id,))
    chunks = conn.execute("""
        SELECT id, title, content
        FROM kb_chunks
        WHERE document_id = ? AND enabled = 1
    """, (document_id,)).fetchall()
    rows = []
    for chunk in chunks:
        counts = Counter()
        metadata = " ".join(str(document[key] or '') for key in (
            'collection', 'title', 'product_name', 'version', 'effective_date',
            'trusted_level', 'dominant_language', 'language_summary'
        ))
        counts.update(index_terms_for_text(metadata, base_weight=3.0))
        counts.update(index_terms_for_text(chunk['title'] or '', base_weight=2.0))
        counts.update(index_terms_for_text(chunk['content'] or '', base_weight=1.0))
        for term, value in counts.most_common(1800):
            if len(term) >= 2 and term not in STOP_TERMS:
                rows.append((term, chunk['id'], document_id, round(float(value), 4)))
    if rows:
        conn.executemany("""
            INSERT OR REPLACE INTO kb_chunk_terms(term, chunk_id, document_id, tf)
            VALUES (?, ?, ?, ?)
        """, rows)
    return len(rows)


def rebuild_all_indexes(db_path):
    init_db(db_path)
    with connect(db_path) as conn:
        conn.execute("DELETE FROM kb_chunk_terms")
        doc_ids = [row['id'] for row in conn.execute("SELECT id FROM kb_documents").fetchall()]
        total_terms = 0
        for doc_id in doc_ids:
            total_terms += build_index_for_document(conn, doc_id)
        return {'documents': len(doc_ids), 'terms': total_terms}


def ensure_index_ready(conn):
    term_count = conn.execute("SELECT COUNT(*) FROM kb_chunk_terms").fetchone()[0]
    chunk_count = conn.execute("SELECT COUNT(*) FROM kb_chunks").fetchone()[0]
    if chunk_count and not term_count:
        doc_ids = [row['id'] for row in conn.execute("SELECT id FROM kb_documents").fetchall()]
        for doc_id in doc_ids:
            build_index_for_document(conn, doc_id)


def list_documents(db_path):
    init_db(db_path)
    with connect(db_path) as conn:
        return [dict(row) for row in conn.execute("""
            SELECT d.*,
                   COALESCE(t.index_terms, 0) AS index_terms
            FROM kb_documents d
            LEFT JOIN (
                SELECT document_id, COUNT(*) AS index_terms
                FROM kb_chunk_terms
                GROUP BY document_id
            ) t ON t.document_id = d.id
            ORDER BY d.updated_at DESC, d.id DESC
        """).fetchall()]


def get_index_stats(db_path):
    init_db(db_path)
    with connect(db_path) as conn:
        return {
            'documents': conn.execute("SELECT COUNT(*) FROM kb_documents").fetchone()[0],
            'chunks': conn.execute("SELECT COUNT(*) FROM kb_chunks").fetchone()[0],
            'terms': conn.execute("SELECT COUNT(*) FROM kb_chunk_terms").fetchone()[0],
            'unique_terms': conn.execute("SELECT COUNT(DISTINCT term) FROM kb_chunk_terms").fetchone()[0],
        }


def query_has_numeric_intent(query):
    text = str(query or '')
    return any(term in text for term in NUMERIC_QUESTION_TERMS) or bool(re.search(r'\d', text))


def term_near_number(text, term, window_chars=90):
    compact = normalize_for_search(text)
    term = normalize_for_search(term)
    if not compact or not term:
        return False
    for match in NUMERIC_PATTERN.finditer(compact):
        start = max(0, match.start() - window_chars)
        end = min(len(compact), match.end() + window_chars)
        if term in compact[start:end]:
            return True
    return False


def score_candidate_features(query, searchable_text, matched_terms, core_terms, numeric_intent):
    compact_variants = text_compact_variants(searchable_text)
    query_compact = normalize_for_search(query)
    score = 0.0
    direct_hits = []

    if 4 <= len(query_compact) <= 36 and term_in_compact_variants(query_compact, compact_variants):
        score += 18.0

    for term in core_terms:
        if term_in_compact_variants(term, compact_variants):
            direct_hits.append(term)
            score += min(9.0, 2.0 + len(term) * 0.8) * query_term_weight(term, query) / 3.0

    unique_core = set(core_terms)
    coverage = (len(set(direct_hits)) / len(unique_core)) if unique_core else 0.0
    score += coverage * 16.0

    year_terms = [term for term in core_terms if YEAR_TERM_PATTERN.match(term)]
    if year_terms:
        if any(term_in_compact_variants(term, compact_variants) for term in year_terms):
            score += 10.0
        else:
            score -= 4.0

    if numeric_intent:
        if NUMERIC_PATTERN.search(searchable_text):
            score += 4.0
        non_numeric_hits = [
            term for term in direct_hits
            if not YEAR_TERM_PATTERN.match(term) and not re.fullmatch(r'\d[\d,.]*', term)
        ]
        if any(term_near_number(searchable_text, term) for term in non_numeric_hits):
            score += 12.0

    return score, coverage, direct_hits


def direct_scan_knowledge(conn, query, top_k=5, min_score=8, candidate_limit=800):
    core_terms = build_core_query_terms(query)
    if not core_terms:
        return []
    numeric_intent = query_has_numeric_intent(query)
    rows = conn.execute("""
        SELECT
            c.id, c.content, c.page_number, c.chunk_index,
            d.title, d.collection, d.product_name, d.version,
            d.effective_date, d.trusted_level
        FROM kb_chunks c
        JOIN kb_documents d ON d.id = c.document_id
        WHERE c.enabled = 1 AND d.enabled = 1
        ORDER BY d.updated_at DESC, c.chunk_index ASC
        LIMIT ?
    """, (max(candidate_limit, top_k * 20),)).fetchall()
    results = []
    for row in rows:
        detail = dict(row)
        searchable_text = "\n".join(str(detail.get(key) or '') for key in (
            'title', 'collection', 'product_name', 'version', 'effective_date', 'trusted_level', 'content'
        ))
        score, coverage, direct_hits = score_candidate_features(
            query, searchable_text, set(), core_terms, numeric_intent
        )
        content = detail.get('content') or ''
        if numeric_intent and NUMERIC_PATTERN.search(content):
            score += 4
        if detail.get('page_number'):
            score += 0.5
        if coverage <= 0 and score <= 0:
            continue
        detail['score'] = round(score, 2)
        detail['coverage'] = round(coverage, 3)
        detail['matched_terms'] = ', '.join(sorted(direct_hits, key=len, reverse=True)[:10])
        detail['snippet'] = make_debug_snippet(content, direct_hits)
        results.append(detail)
    results.sort(key=lambda item: item['score'], reverse=True)
    if min_score is not None:
        results = [item for item in results if item.get('score', 0) >= float(min_score)]
    return results[:top_k]


def make_debug_snippet(content, matched_terms, max_chars=420):
    text = str(content or '').strip()
    if len(text) <= max_chars:
        return text
    best_pos = -1
    preferred_terms = [
        term for term in (matched_terms or [])
        if not re.fullmatch(r'\d{4}\u5e74|\d{4}|\d[\d,.]*', str(term or ''))
    ]
    compact_text = normalize_for_search(text)
    normalized_preferred = [normalize_for_search(term) for term in preferred_terms if len(str(term or '')) >= 2]
    for number_match in NUMERIC_PATTERN.finditer(compact_text):
        start = max(0, number_match.start() - 90)
        end = min(len(compact_text), number_match.end() + 90)
        window = compact_text[start:end]
        if any(term and term in window for term in normalized_preferred):
            raw_positions = [text.find(term) for term in preferred_terms if len(str(term or '')) >= 2 and text.find(term) >= 0]
            best_pos = min(raw_positions) if raw_positions else number_match.start()
            break
    if best_pos < 0:
        for term in sorted(preferred_terms, key=len, reverse=True):
            if len(term) < 2:
                continue
            pos = text.find(term)
            if pos >= 0:
                best_pos = pos
                break
    if best_pos < 0:
        for term in sorted(matched_terms or [], key=len, reverse=True):
            if len(term) < 2:
                continue
            pos = text.find(term)
            if pos >= 0:
                best_pos = pos
                break
    if best_pos < 0:
        number_match = NUMERIC_PATTERN.search(text)
        if number_match:
            best_pos = number_match.start()
    if best_pos < 0:
        return text[:max_chars].rstrip() + '...'
    start = max(0, best_pos - max_chars // 14)
    end = min(len(text), start + max_chars)
    start = max(0, end - max_chars)
    prefix = '...' if start > 0 else ''
    suffix = '...' if end < len(text) else ''
    return prefix + text[start:end].strip() + suffix


def search_knowledge(db_path, query, top_k=5, min_score=8, candidate_limit=240, auto_rebuild=True):
    init_db(db_path)
    terms = build_query_terms(query)
    term_weights = {term: query_term_weight(term, query) for term in terms}
    core_terms = build_core_query_terms(query)
    if not terms:
        return []
    with connect(db_path) as conn:
        if auto_rebuild:
            ensure_index_ready(conn)
        placeholders = ','.join(['?'] * len(terms))
        rows = conn.execute(f"""
            SELECT term, chunk_id, document_id, tf
            FROM kb_chunk_terms
            WHERE term IN ({placeholders})
        """, terms).fetchall()
        if not rows:
            return direct_scan_knowledge(
                conn,
                query,
                top_k=top_k,
                min_score=min_score,
                candidate_limit=candidate_limit
            )
        total_chunks = conn.execute("SELECT COUNT(*) FROM kb_chunks WHERE enabled = 1").fetchone()[0] or 1
        df = Counter(row['term'] for row in rows)
        per_chunk = defaultdict(lambda: {'score': 0.0, 'terms': set(), 'document_id': None})
        for row in rows:
            term = row['term']
            idf = math.log((total_chunks + 1) / (df[term] + 1)) + 1.0
            length_boost = min(3.0, max(1.0, len(term) / 3.0))
            rare_boost = 2.5 if df[term] <= max(3, total_chunks // 50) else 1.0
            score = float(row['tf']) * idf * length_boost * rare_boost * term_weights.get(term, 1.0)
            item = per_chunk[row['chunk_id']]
            item['score'] += score
            item['terms'].add(term)
            item['document_id'] = row['document_id']
        ranked = sorted(per_chunk.items(), key=lambda item: item[1]['score'], reverse=True)[:candidate_limit]
        chunk_ids = [item[0] for item in ranked]
        placeholders = ','.join(['?'] * len(chunk_ids))
        details = {
            row['id']: dict(row)
            for row in conn.execute(f"""
                SELECT
                    c.id, c.content, c.page_number, c.chunk_index,
                    d.title, d.collection, d.product_name, d.version,
                    d.effective_date, d.trusted_level
                FROM kb_chunks c
                JOIN kb_documents d ON d.id = c.document_id
                WHERE c.id IN ({placeholders}) AND c.enabled = 1 AND d.enabled = 1
            """, chunk_ids).fetchall()
        }
    numeric_intent = query_has_numeric_intent(query)
    results = []
    for chunk_id, scoring in ranked:
        detail = details.get(chunk_id)
        if not detail:
            continue
        score = scoring['score']
        content = detail.get('content') or ''
        searchable_text = "\n".join(str(detail.get(key) or '') for key in (
            'title', 'collection', 'product_name', 'version', 'effective_date', 'trusted_level', 'content'
        ))
        feature_score, coverage, direct_hits = score_candidate_features(
            query, searchable_text, scoring['terms'], core_terms, numeric_intent
        )
        score += feature_score
        scoring['terms'].update(direct_hits)
        if core_terms and coverage < 0.12:
            score *= 0.65
        if numeric_intent and NUMERIC_PATTERN.search(content):
            score += 8
        if detail.get('page_number'):
            score += 0.5
        detail['score'] = round(score, 2)
        detail['coverage'] = round(coverage, 3)
        detail['matched_terms'] = ', '.join(sorted(scoring['terms'], key=len, reverse=True)[:10])
        snippet_terms = [
            term for term in direct_hits
            if len(term) >= 2
            and not YEAR_TERM_PATTERN.match(term)
            and not re.fullmatch(r'\d[\d,.]*', term)
            and term not in STOP_TERMS
            and not any(stop in term for stop in STOP_TERMS if len(stop) >= 2 and re.search(r'[\u4e00-\u9fff]', stop))
        ]
        detail['snippet'] = make_debug_snippet(content, snippet_terms or scoring['terms'])
        results.append(detail)
    results.sort(key=lambda item: item['score'], reverse=True)
    if min_score is not None:
        results = [item for item in results if item.get('score', 0) >= float(min_score)]
    return results[:top_k]


def format_knowledge_context(chunks, max_chars=2500):
    lines = []
    used_chars = 0
    for index, item in enumerate(chunks, start=1):
        source_parts = [item.get('title') or '未命名资料']
        if item.get('collection'):
            source_parts.append(item['collection'])
        if item.get('product_name'):
            source_parts.append(item['product_name'])
        if item.get('version'):
            source_parts.append(f"版本:{item['version']}")
        if item.get('effective_date'):
            source_parts.append(f"生效:{item['effective_date']}")
        if item.get('page_number'):
            source_parts.append(f"第{item['page_number']}页")
        content = str(item.get('content') or '').strip()
        entry = f"[资料{index}] {' | '.join(source_parts)}\n{content}\n"
        remaining = max_chars - used_chars
        if remaining <= 0:
            break
        if len(entry) > remaining:
            entry = entry[:remaining].rstrip()
        lines.append(entry)
        used_chars += len(entry)
    return "\n".join(lines).strip()
